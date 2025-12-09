import streamlit as st
import pandas as pd
import numpy as np
from datetime import date, timedelta
import plotly.express as px
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURA√á√ÉO DA P√ÅGINA ---
st.set_page_config(page_title="Gest√£o AIA - Pro", layout="wide", page_icon="‚öñÔ∏è")

# --- 1. BASE DE DADOS DE FERIADOS (VALIDADA) ---
feriados_nacionais = [
    "2025-01-01", "2025-04-18", "2025-04-20", "2025-04-25", "2025-05-01",
    "2025-06-10", "2025-06-19", "2025-08-15", "2025-10-05", "2025-11-01",
    "2025-12-01", "2025-12-08", "2025-12-25",
    "2026-01-01", "2026-04-03", "2026-04-05", "2026-04-25", "2026-05-01",
    "2026-06-04", "2026-06-10", "2026-08-15", "2026-10-05", "2026-11-01",
    "2026-12-01", "2026-12-08", "2026-12-25"
]
feriados_np = np.array(feriados_nacionais, dtype='datetime64[D]')

# --- 2. FUN√á√ïES DE C√ÅLCULO ---
def somar_dias_uteis(data_inicio, dias, feriados):
    """Calcula data futura somando dias √∫teis."""
    return np.busday_offset(np.datetime64(data_inicio), dias, roll='forward', weekmask='1111100', holidays=feriados)

def formatar_data(np_date):
    """Formata data para PT."""
    return pd.to_datetime(np_date).strftime("%d/%m/%Y")

# --- 3. GERADOR DE RELAT√ìRIO WORD (JURIDICAMENTE ATUALIZADO) ---
def gerar_relatorio_completo(df_dados, data_fim, prazo_max, saldo, fig_timeline):
    doc = Document()
    
    # Cabe√ßalho
    titulo = doc.add_heading('Cronograma de Prazos AIA', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Data de Emiss√£o: {date.today().strftime("%d/%m/%Y")}')
    doc.add_paragraph('')

    # 1. Enquadramento Legal (ATUALIZADO)
    doc.add_heading('1. Enquadramento Legal', level=1)
    
    texto_legal = (
        "A presente calendariza√ß√£o foi elaborada nos termos do Regime Jur√≠dico da Avalia√ß√£o de Impacte Ambiental (RJAIA), "
        "aprovado pelo Decreto-Lei n.¬∫ 151-B/2013, conjugado com o C√≥digo do Procedimento Administrativo (CPA).\n"
    )
    p = doc.add_paragraph(texto_legal)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Lista de fundamentos
    p_details = doc.add_paragraph()
    p_details.add_run("1. Contagem: ").bold = True
    p_details.add_run(
        "Os prazos administrativos contam-se em dias √∫teis (Art. 87.¬∫ do CPA), suspendendo-se aos s√°bados, domingos e feriados nacionais. "
        "N√£o h√° suspens√£o durante f√©rias judiciais.\n"
    )
    p_details.add_run("2. Suspens√µes: ").bold = True
    p_details.add_run(
        "O prazo de decis√£o suspende-se sempre que a Autoridade aguarde elementos do proponente. "
        "Esta suspens√£o fundamenta-se no "
    )
    p_details.add_run("Art. 13.¬∫, n.¬∫ 4 do RJAIA ").bold = True
    p_details.add_run("(fase de conformidade/aperfei√ßoamento) e no ")
    p_details.add_run("Art. 16.¬∫ do RJAIA ").bold = True
    p_details.add_run("(fase de an√°lise t√©cnica), em articula√ß√£o com o princ√≠pio geral do ")
    p_details.add_run("Art. 117.¬∫, n.¬∫ 2 do CPA.").bold = True

    # 2. Resumo Executivo
    doc.add_heading('2. Resumo de Prazos', level=1)
    
    p_resumo = doc.add_paragraph()
    run_dt = p_resumo.add_run(f'Data Limite da Decis√£o (DIA): {data_fim}')
    run_dt.bold = True
    run_dt.font.size = Pt(12)
    
    doc.add_paragraph(f'Prazo Legal Total: {prazo_max} dias √∫teis')
    
    if saldo >= 0:
        doc.add_paragraph(f'Saldo Dispon√≠vel: {saldo} dias √∫teis')
    else:
        p_alert = doc.add_paragraph()
        r_alert = p_alert.add_run(f'DERRAPAGEM: {abs(saldo)} dias acima do prazo.')
        r_alert.bold = True
        r_alert.font.color.rgb = None 

    # 3. Infograma
    doc.add_heading('3. Linha do Tempo Visual', level=1)
    try:
        img_buffer = BytesIO()
        # Nota: Requer kaleido==0.2.1 no requirements.txt
        fig_timeline.write_image(img_buffer, format='png', width=800, height=400)
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph("[Gr√°fico indispon√≠vel nesta vers√£o. Verifique biblioteca 'kaleido']")

    # 4. Tabela
    doc.add_page_break()
    doc.add_heading('4. Tabela Detalhada das Etapas', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Fase'
    hdr[1].text = 'Dura√ß√£o'
    hdr[2].text = 'In√≠cio'
    hdr[3].text = 'Fim'

    for _, row in df_dados.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Fase'])
        cells[1].text = str(row['Dura√ß√£o'])
        cells[2].text = str(row['In√≠cio'])
        cells[3].text = str(row['Fim'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. INTERFACE STREAMLIT ---
st.title("üìÖ Gest√£o de Prazos AIA")
st.markdown("Simulador de Prazos RJAIA/CPA com gest√£o de suspens√µes e relat√≥rios.")

with st.sidebar:
    st.header("1. Configura√ß√£o Base")
    tipo = st.radio("Tipologia:", ["AIA Geral (150 dias)", "AIA Simplificado (90 dias)"])
    prazo_max = 150 if "Geral" in tipo else 90
    data_inicio = st.date_input("Data de Submiss√£o", date(2025, 6, 3))
    
    st.markdown("---")
    st.header("2. Fases e Suspens√µes")
    
    # FASE 1
    st.subheader("Fase 1: Conformidade")
    d1 = st.number_input("Dura√ß√£o (Dias √öteis)", 10, key="d1")
    susp_conf = st.number_input("Suspens√£o / Aperfei√ßoamento (Dias Corridos)", value=0, help="Art. 13¬∫ RJAIA: Convite ao aperfei√ßoamento.", key="s1")
    
    # FASE 2
    st.subheader("Fase 2: Consulta P√∫blica")
    d2 = st.number_input("Dura√ß√£o (Dias √öteis)", 30, key="d2")
    
    # FASE 3
    st.subheader("Fase 3: An√°lise T√©cnica")
    d3 = st.number_input("Dura√ß√£o (Dias √öteis)", 60, key="d3")
    susp_adit = st.number_input("Suspens√£o / Aditamentos (Dias Corridos)", value=45, help="Art. 16¬∫ RJAIA: Pedido de elementos adicionais.", key="s3")
    
    # FASE 4
    st.subheader("Fase 4: Audi√™ncia Pr√©via")
    d4 = st.number_input("Dura√ß√£o (Dias √öteis)", 10, key="d4")
    susp_aud = st.number_input("Suspens√£o da Contagem (Dias √öteis)", value=10, help="Art. 117¬∫ CPA: Suspens√£o para pron√∫ncia.", key="s4")
    
    # FASE 5
    st.subheader("Fase 5: Decis√£o (DIA)")
    dias_restantes_calc = prazo_max - (d1+d2+d3+d4)
    d5 = st.number_input("Dura√ß√£o Restante (Dias √öteis)", value=dias_restantes_calc, disabled=True)

# --- 5. MOTOR DE C√ÅLCULO ---
cronograma = []
cursor = data_inicio
dias_consumidos = 0

# --- L√≥gica Passo a Passo ---

# 1. CONFORMIDADE
inicio = cursor
fim_np = somar_dias_uteis(inicio, d1, feriados_np)
fim = pd.to_datetime(fim_np).date()
cronograma.append({"Fase": "1. Conformidade", "In√≠cio": formatar_data(inicio), "Fim": formatar_data(fim), "Start": inicio, "Finish": fim, "Dura√ß√£o": f"{d1} √∫teis", "Tipo": "Consome Prazo"})
cursor = fim
dias_consumidos += d1

# Suspens√£o Conformidade (NOVO)
if susp_conf > 0:
    inicio_susp = cursor
    fim_susp = cursor + timedelta(days=susp_conf) # Dias Corridos
    cronograma.append({"Fase": "‚ö†Ô∏è Aperfei√ßoamento (Art. 13¬∫)", "In√≠cio": formatar_data(inicio_susp), "Fim": formatar_data(fim_susp), "Start": inicio_susp, "Finish": fim_susp, "Dura√ß√£o": f"{susp_conf} corridos", "Tipo": "Suspens√£o"})
    cursor = fim_susp

# 2. CONSULTA P√öBLICA
inicio = cursor
fim_np = somar_dias_uteis(inicio, d2, feriados_np)
fim = pd.to_datetime(fim_np).date()
cronograma.append({"Fase": "2. Consulta P√∫blica", "In√≠cio": formatar_data(inicio), "Fim": formatar_data(fim), "Start": inicio, "Finish": fim, "Dura√ß√£o": f"{d2} √∫teis", "Tipo": "Consome Prazo"})
cursor = fim
dias_consumidos += d2

# 3. AN√ÅLISE T√âCNICA
inicio = cursor
fim_np = somar_dias_uteis(inicio, d3, feriados_np)
fim = pd.to_datetime(fim_np).date()
cronograma.append({"Fase": "3. An√°lise T√©cnica", "In√≠cio": formatar_data(inicio), "Fim": formatar_data(fim), "Start": inicio, "Finish": fim, "Dura√ß√£o": f"{d3} √∫teis", "Tipo": "Consome Prazo"})
cursor = fim
dias_consumidos += d3

# Suspens√£o Aditamentos
if susp_adit > 0:
    inicio_susp = cursor
    fim_susp = cursor + timedelta(days=susp_adit) # Dias Corridos
    cronograma.append({"Fase": "‚è∏Ô∏è Aditamentos (Art. 16¬∫)", "In√≠cio": formatar_data(inicio_susp), "Fim": formatar_data(fim_susp), "Start": inicio_susp, "Finish": fim_susp, "Dura√ß√£o": f"{susp_adit} corridos", "Tipo": "Suspens√£o"})
    cursor = fim_susp

# 4. AUDI√äNCIA PR√âVIA
# Ajuste: Se suspens√£o acabou ao FDS, come√ßar em dia √∫til
cursor_util = pd.to_datetime(somar_dias_uteis(cursor, 0, feriados_np)).date()
inicio = cursor_util
fim_np = somar_dias_uteis(inicio, d4, feriados_np)
fim = pd.to_datetime(fim_np).date()
cronograma.append({"Fase": "4. Audi√™ncia Pr√©via", "In√≠cio": formatar_data(inicio), "Fim": formatar_data(fim), "Start": inicio, "Finish": fim, "Dura√ß√£o": f"{d4} √∫teis", "Tipo": "Consome Prazo"})
cursor = fim
dias_consumidos += d4

# Suspens√£o Audi√™ncia (Dias √öteis - Art 117 CPA)
if susp_aud > 0:
    inicio_susp = cursor
    fim_susp_np = somar_dias_uteis(inicio_susp, susp_aud, feriados_np) # Dias √öteis
    fim_susp = pd.to_datetime(fim_susp_np).date()
    cronograma.append({"Fase": "‚è∏Ô∏è An√°lise Pron√∫ncias (CPA)", "In√≠cio": formatar_data(inicio_susp), "Fim": formatar_data(fim_susp), "Start": inicio_susp, "Finish": fim_susp, "Dura√ß√£o": f"{susp_aud} √∫teis", "Tipo": "Suspens√£o"})
    cursor = fim_susp

# 5. DECIS√ÉO FINAL (Saldo)
dias_finais = prazo_max - dias_consumidos
if dias_finais > 0:
    inicio = cursor
    fim_np = somar_dias_uteis(inicio, dias_finais, feriados_np)
    fim = pd.to_datetime(fim_np).date()
    cronograma.append({"Fase": "5. Emiss√£o da DIA", "In√≠cio": formatar_data(inicio), "Fim": formatar_data(fim), "Start": inicio, "Finish": fim, "Dura√ß√£o": f"{dias_finais} √∫teis", "Tipo": "Consome Prazo"})
    cursor = fim
    dias_consumidos += dias_finais

df = pd.DataFrame(cronograma)
data_final_txt = formatar_data(cursor)
saldo = prazo_max - dias_consumidos

# --- 6. VISUALIZA√á√ÉO ---
st.divider()

col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("Linha do Tempo (Gantt)")
    fig = px.timeline(
        df, x_start="Start", x_end="Finish", y="Fase", color="Tipo",
        color_discrete_map={"Consome Prazo": "#2E86C1", "Suspens√£o": "#E74C3C"},
        hover_data=["Dura√ß√£o"],
        title=f"Previs√£o de Fim: {data_final_txt}"
    )
    fig.update_yaxes(autorange="reversed")
    fig.update_layout(height=400)
    st.plotly_chart(fig, use_container_width=True)

with col2:
    st.subheader("Resumo Oficial")
    st.metric("Data Final (DIA)", data_final_txt)
    st.metric("Dias Consumidos", f"{dias_consumidos} / {prazo_max}")
    
    st.write("---")
    st.write("üìÑ **Documenta√ß√£o**")
    
    # Tratamento de erro caso kaleido falhe na primeira execu√ß√£o
    try:
        arquivo = gerar_relatorio_completo(df, data_final_txt, prazo_max, saldo, fig)
        st.download_button(
            "üì• Baixar Relat√≥rio (.docx)",
            data=arquivo,
            file_name=f"Cronograma_AIA_{date.today()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Erro ao gerar relat√≥rio: {e}")
        st.caption("Verifique se 'kaleido==0.2.1' est√° no requirements.txt")

st.divider()
with st.expander("Ver Tabela de Dados Completa"):
    st.dataframe(df[['Fase', 'In√≠cio', 'Fim', 'Dura√ß√£o', 'Tipo']], use_container_width=True)
